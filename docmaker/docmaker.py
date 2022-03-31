import urllib.request, urllib.parse, urllib.error
import json
import ssl
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.dml import MSO_THEME_COLOR_INDEX
import docx
import re
from datetime import datetime, timedelta

def date_string():
    dt = datetime.today()
    day = dt.weekday()

    print(dt,day)
    print(type(dt.date()))

    if not day == 0:
        date_str = str(dt.date() - timedelta(days=1))
    else:
        date_str = str(dt.date() - timedelta(days=2))

    print(date_str)

    return date_str

def add_content(paragraph, text):
            if  text:

                cstring = re.findall('.+- (.+)\.+', text)

                if not len(cstring) > 0:
                    cstring = (re.findall('\) (.+)\.+', text))

                if not len(cstring) > 0:
                    cstring = (re.findall('(.+)\[', text))

                paragraph.add_run(cstring)

            else:
                paragraph.add_run('NO CONTENT FOUND')

def add_hyperlink(paragraph, text, url):
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run ()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def make_url(keyterm, date):
    #Format
    keyterm = 'q=' + keyterm + '&'

    url = ('http://newsapi.org/v2/everything?'
            #'country=us&'
            + keyterm +
            'excludedomains='
                'economictimes.indiatimes.com,'
                'thehindu.com,'
                'indianexpress.com,'
                'timesofindia.com,'
                'indiatimes.com,'
                'aljazeera.com,'
                'bbc.co.uk,'
                'theguardian.com,'
                'theguardian.comworld,'
                'thenextweb.com,'
                'voxeu.org,'
                'theregister.co.uk,'
                'searchenginejournal.com,'
                'medium.com,'
                'telegraph.co.uk,'
                'eater.com,'
                'maketscreener.com,'
                'mashable.com,'
                'cjr.org,'
                'independant.co.uk,'
                'epsncricinfo.com,'
                'moneycontrol.com,'
                'rt.com,'
                'viewfromthewing.com'
                    '&'
            'pageSize=50&'
            'from='
                + date +
                '&'
            'language=en&'
            'sortby=popularity&'
            'apiKey=43bb493c38b24d6a95556af541a953a6')

    #print(url)

    return(url)

def jget_news_data(url):
    data = urllib.request.urlopen(url)
    js = json.load(data)
    print(js.keys())

    return(js)

def make_section(document,keyterm,date,spam,titles,conswitch,):

#####Intitialize Variables
    sourced = {}

#####Create URl and Load into JSON OBJECT
    url = make_url(keyterm, date)
    js = jget_news_data(url)
    dup_limit = 0


#####Loop Through Article
    for i in js['articles']:


        #Set Variables
        title = i['title']
        source_name = i['source']['name']
        author = i['author']

        #Control Source Sites
        #if source_name in blacklist:
            #continue


        #Check Duplicate Titles
        if (title in titles) and (dup_limit == 1):
            #continue
            titles.append(title)
        else:
            titles.append(title)


        #Limit Source Spam
        if source_name in sourced:
            sourced[source_name] = sourced[source_name] + 1
        else:
            sourced[source_name] = 1

        if sourced[source_name] > spam:
            continue


        #Populate Document
        p = document.add_paragraph('', style='List Bullet')


        #Add Source
        p.add_run(source_name).bold = True
        p.add_run(', ')


        #Add Link
        add_hyperlink(p, title, i['url'])
        p.add_run(', by ')


        #Add Author
        if author != None:
            if  source_name != 'Associated Press':
                author = author.lower()
                author = author.title()
                p.add_run(author)
            else:
                author = author.lower()
                author = author.title()
                author_str = re.findall('By (.+) Associated Press',author)
                if len(author_str) == 0:
                    author_str = re.findall('By (.+) Ap',author)
                p.add_run(author_str)
        else:
            p.add_run(' NO AUTHOR FOUND ')

        #Add Content
        if not conswitch == False:
            p.add_run(': "')
            add_content(p,i['content'])
            p.add_run('"')

        #Adds Space Between Entires
        p.add_run('\n')


    return document, titles, js
